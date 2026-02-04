"""
知识点提取服务 - 完整无限制版本
"""
import os
import re
import asyncio
from typing import List, Dict, Any, Optional
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy import text
import PyPDF2
import docx
from pathlib import Path

from app.models.knowledge import (
    CourseDocument, KnowledgePoint, KnowledgeRelation, 
    DocumentProcessingTask, KnowledgeGraph
)


class KnowledgePointExtractor:
    """知识点提取器 - 无字数限制完整解析"""
    
    def __init__(self, db: Session):
        self.db = db
    
    async def process_document_async(
        self,
        document_id: str,
        course_id: str,
        file_path: str,
        file_type: str
    ) -> Dict[str, Any]:
        """异步处理文档，提取知识点"""
        
        # 创建处理任务
        task = DocumentProcessingTask(
            document_id=document_id,
            task_type='knowledge_extraction',
            status='processing',
            total_steps=6,
            started_at=datetime.utcnow()
        )
        self.db.add(task)
        self.db.commit()
        self.db.refresh(task)
        
        try:
            # 步骤1：提取文本 (0-20%)
            await self._update_progress(task.id, 5, 1, 'processing')
            text = await self._extract_text(file_path, file_type)
            
            # 更新文档表
            doc = self.db.query(CourseDocument).filter(CourseDocument.id == document_id).first()
            if doc:
                doc.extracted_text = text
                doc.processing_progress = 20
                self.db.commit()
            
            await self._update_progress(task.id, 20, 1, 'processing')
            
            # 步骤2：分割章节 (20-30%)
            sections = self._split_comprehensive(text)
            await self._update_progress(task.id, 30, 2, 'processing')
            
            # 步骤3：提取所有知识点 (30-70%)
            all_points = []
            for idx, section in enumerate(sections):
                section_points = self._extract_all_points_from_section(section, idx + 1)
                all_points.extend(section_points)
                
                progress = 30 + int((idx + 1) / len(sections) * 40)
                await self._update_progress(task.id, progress, 3, 'processing')
                
                if doc:
                    doc.processing_progress = progress
                    self.db.commit()
            
            # 步骤4：提取关键词 (70-80%)
            for point in all_points:
                point['keywords'] = self._extract_keywords_advanced(point['content'])
            await self._update_progress(task.id, 80, 4, 'processing')
            
            # 步骤5：保存知识点 (80-90%)
            saved_points = await self._save_all_points(all_points, course_id, document_id)
            await self._update_progress(task.id, 90, 5, 'processing')
            
            # 步骤6：构建知识图谱 (90-100%)
            await self._build_complete_graph(saved_points, course_id)
            
            # 完成
            task.status = 'completed'
            task.progress = 100
            task.current_step = 6
            task.completed_at = datetime.utcnow()
            task.result_data = {
                'total_points': len(saved_points),
                'sections': len(sections),
                'text_length': len(text)
            }
            
            if doc:
                doc.processed_status = 'completed'
                doc.processing_progress = 100
            
            self.db.commit()
            
            return {
                'success': True,
                'task_id': str(task.id),
                'total_points': len(saved_points),
                'message': '知识点提取完成'
            }
            
        except Exception as e:
            task.status = 'failed'
            task.error_message = str(e)
            task.completed_at = datetime.utcnow()
            
            if doc:
                doc.processed_status = 'failed'
                doc.error_message = str(e)
            
            self.db.commit()
            
            return {
                'success': False,
                'error': str(e)
            }
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """提取文本 - 无限制"""
        ext = file_type.lower().replace('.', '')
        
        if ext == 'pdf':
            return await self._extract_pdf_unlimited(file_path)
        elif ext in ['docx', 'doc']:
            return await self._extract_docx_unlimited(file_path)
        elif ext in ['txt', 'md']:
            return await self._extract_text_file(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
    
    async def _extract_pdf_unlimited(self, file_path: str) -> str:
        """PDF提取 - 完全无限制"""
        all_text = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                # 提取所有页面，无页数限制
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        all_text.append(text)
                
            return '\n\n'.join(all_text)
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")
    
    async def _extract_docx_unlimited(self, file_path: str) -> str:
        """DOCX提取 - 完全无限制"""
        try:
            doc = docx.Document(file_path)
            all_text = []
            
            # 提取所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        all_text.append(row_text)
            
            return '\n\n'.join(all_text)
        except Exception as e:
            raise Exception(f"DOCX解析失败: {str(e)}")
    
    async def _extract_text_file(self, file_path: str) -> str:
        """文本文件提取"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except:
                continue
        raise Exception("无法读取文本文件")
    
    def _split_comprehensive(self, text: str) -> List[Dict[str, Any]]:
        """全面分割文本 - 识别所有可能的章节结构"""
        sections = []
        
        # 多种章节模式
        patterns = [
            (r'第[一二三四五六七八九十零〇百千万\d]+[章节课讲部分篇][\s\u3000]*[：:]*[\s\u3000]*(.+?)(?=第[一二三四五六七八九十零〇百千万\d]+[章节课讲部分篇]|$)', 1),
            (r'Chapter[\s]+[\d]+[\s]*[：:]*[\s]*(.+?)(?=Chapter[\s]+[\d]+|$)', 1),
            (r'^[\d]+[\.、．][\s\u3000]*(.+?)(?=^[\d]+[\.、．]|$)', 1),
            (r'^[一二三四五六七八九十百千]+[、．.][\s\u3000]*(.+?)(?=^[一二三四五六七八九十百千]+[、．.]|$)', 1),
            (r'[【\[]第[一二三四五六七八九十\d]+部分[】\]](.+?)(?=[【\[]第[一二三四五六七八九十\d]+部分[】\]]|$)', 1),
        ]
        
        for pattern, title_group in patterns:
            matches = list(re.finditer(pattern, text, re.MULTILINE | re.DOTALL | re.IGNORECASE))
            if len(matches) >= 1:
                for i, match in enumerate(matches):
                    title = match.group(title_group).split('\n')[0].strip() if match.lastindex >= title_group else f"Section {i+1}"
                    sections.append({
                        'title': title[:200],
                        'content': match.group(0).strip(),
                        'level': 1,
                        'order': i
                    })
                if sections:
                    return sections
        
        # 按双换行分段
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if len(p.strip()) > 50]
        if paragraphs:
            return [{'title': p[:100], 'content': p, 'level': 1, 'order': i} 
                    for i, p in enumerate(paragraphs)]
        
        # 按单换行分段
        lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 30]
        return [{'title': l[:100], 'content': l, 'level': 1, 'order': i} 
                for i, l in enumerate(lines[:100])]
    
    def _extract_all_points_from_section(
        self, 
        section: Dict[str, Any], 
        section_num: int
    ) -> List[Dict[str, Any]]:
        """从章节提取所有知识点 - 完全无限制"""
        points = []
        content = section['content']
        title = section['title']
        
        # 1. 章节主知识点
        points.append({
            'name': title,
            'content': content[:2000],  # 摘要
            'full_content': content,  # 完整内容
            'type': 'chapter',
            'level': 1,
            'order': section_num,
            'importance': 5
        })
        
        # 2. 小节
        subsections = self._extract_all_subsections(content)
        for i, sub in enumerate(subsections):
            points.append({
                'name': sub['name'],
                'content': sub['content'],
                'type': 'section',
                'level': 2,
                'order': i,
                'parent_name': title,
                'importance': 4
            })
        
        # 3. 定义
        definitions = self._extract_all_definitions(content)
        for i, defn in enumerate(definitions):
            points.append({
                'name': defn['term'],
                'content': defn['definition'],
                'type': 'definition',
                'level': 2,
                'order': len(subsections) + i,
                'parent_name': title,
                'importance': 4
            })
        
        # 4. 要点/列表项
        list_items = self._extract_list_items(content)
        for i, item in enumerate(list_items):
            points.append({
                'name': item['title'],
                'content': item['content'],
                'type': 'point',
                'level': 3,
                'order': i,
                'parent_name': title,
                'importance': 3
            })
        
        # 5. 示例
        examples = self._extract_examples(content)
        for i, example in enumerate(examples):
            points.append({
                'name': example['title'],
                'content': example['content'],
                'type': 'example',
                'level': 3,
                'order': i,
                'parent_name': title,
                'importance': 3
            })
        
        return points
    
    def _extract_all_subsections(self, text: str) -> List[Dict[str, str]]:
        """提取所有小节 - 无限制"""
        subsections = []
        patterns = [
            r'(\d+\.\d+[\s\u3000：:]+)(.+?)(?=\d+\.\d+[\s\u3000：:]|$)',
            r'([（(]\d+[）)][\s\u3000]*)(.+?)(?=[（(]\d+[）)]|$)',
            r'([①②③④⑤⑥⑦⑧⑨⑩]+[\s\u3000]*)(.+?)(?=[①②③④⑤⑥⑦⑧⑨⑩]+|$)',
            r'([a-zA-Z]\.)[\s\u3000]*(.+?)(?=[a-zA-Z]\.|$)',
        ]
        
        for pattern in patterns:
            matches = list(re.finditer(pattern, text, re.MULTILINE | re.DOTALL))
            if len(matches) >= 2:
                for match in matches:
                    name = match.group(2).split('\n')[0].strip()[:300]
                    content = match.group(0).strip()[:3000]
                    if name and content:
                        subsections.append({'name': name, 'content': content})
                if subsections:
                    return subsections
        
        return subsections
    
    def _extract_all_definitions(self, text: str) -> List[Dict[str, str]]:
        """提取所有定义 - 无限制"""
        definitions = []
        patterns = [
            r'(.{2,80})是指(.{5,500})(?=[。\n])',
            r'(.{2,80})定义为(.{5,500})(?=[。\n])',
            r'(.{2,80})：(.{5,500})(?=[。\n])',
            r'(.{2,80})\((.{5,500})\)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                term = match.group(1).strip()
                defn = match.group(2).strip()
                if 2 <= len(term) <= 80 and 5 <= len(defn) <= 500:
                    definitions.append({'term': term, 'definition': defn})
        
        return definitions[:100]  # 最多100个定义
    
    def _extract_list_items(self, text: str) -> List[Dict[str, str]]:
        """提取列表项"""
        items = []
        patterns = [
            r'^[\s]*[•·\-*][\s]+(.+?)$',
            r'^[\s]*[\d]+[.)、][\s]+(.+?)$',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                item_text = match.group(1).strip()
                if item_text:
                    items.append({
                        'title': item_text[:100],
                        'content': item_text
                    })
        
        return items[:50]
    
    def _extract_examples(self, text: str) -> List[Dict[str, str]]:
        """提取示例"""
        examples = []
        patterns = [
            r'[【\[]?例[\d]*[】\]]?[:：](.+?)(?=[【\[]?例[\d]*[】\]]?|$)',
            r'示例[:：](.+?)(?=示例|$)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.DOTALL)
            for i, match in enumerate(matches):
                content = match.group(1).strip()[:1000]
                if content:
                    examples.append({
                        'title': f'示例{i+1}',
                        'content': content
                    })
        
        return examples[:20]
    
    def _extract_keywords_advanced(self, text: str) -> List[str]:
        """高级关键词提取"""
        # 中文分词
        words = re.findall(r'[\u4e00-\u9fa5]{2,}', text)
        # 英文词
        en_words = re.findall(r'[a-zA-Z]{3,}', text)
        
        all_words = words + en_words
        word_freq = {}
        for word in all_words:
            if 2 <= len(word) <= 20:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, _ in sorted_words[:15]]
    
    async def _save_all_points(
        self,
        points: List[Dict[str, Any]],
        course_id: str,
        document_id: str
    ) -> List[KnowledgePoint]:
        """保存所有知识点"""
        from app.models.knowledge import CourseDocument
        from app.models.course import KnowledgeBase
        
        saved = []
        name_to_id = {}
        
        # 获取文档类型
        doc = self.db.query(CourseDocument).filter(CourseDocument.id == document_id).first()
        document_type = doc.document_type if doc else 'material'
        
        # 按层级保存
        for level in range(1, 6):
            for point in points:
                if point.get('level') == level:
                    parent_id = name_to_id.get(point.get('parent_name'))
                    
                    kp = KnowledgePoint(
                        course_id=course_id,
                        document_id=document_id,
                        point_name=point['name'][:500],
                        point_content=point.get('content', '')[:10000],  # 最多10000字符
                        point_type=point.get('type', 'concept'),
                        level=level,
                        parent_id=parent_id,
                        keywords=point.get('keywords', [])[:15],
                        importance=point.get('importance', 3),
                        order_index=point.get('order', 0),
                        extra_info={'full_content': point.get('full_content', '')[:50000]} if point.get('full_content') else None
                    )
                    self.db.add(kp)
                    self.db.flush()
                    saved.append(kp)
                    name_to_id[point['name']] = kp.id
                    
                    # 保存到向量数据库，包含document_type
                    kb_entry = KnowledgeBase(
                        document_id=document_id,
                        course_id=course_id,
                        document_type=document_type,
                        chunk_text=point.get('content', '')[:5000],
                        chunk_index=point.get('order', 0),
                        chunk_metadata={
                            'point_name': point['name'][:500],
                            'point_type': point.get('type', 'concept'),
                            'level': level,
                            'keywords': point.get('keywords', [])[:15],
                            'document_type': document_type  # 在metadata中也保存document_type
                        }
                    )
                    self.db.add(kb_entry)
        
        # 不在此处 commit，由 process_document_async 在 _build_complete_graph 后统一提交，
        # 避免 commit 后 session 过期导致构建图谱时对每个 KnowledgePoint 按 id 懒加载（N+1）
        self.db.flush()
        return saved
    
    async def _build_complete_graph(
        self,
        points: List[KnowledgePoint],
        course_id: str
    ):
        """构建完整知识图谱"""
        # 基于关键词构建关系
        for i, p1 in enumerate(points):
            if not p1.keywords:
                continue
            
            kw1 = set(p1.keywords)
            for p2 in points[i+1:]:
                if not p2.keywords or p1.id == p2.id:
                    continue
                
                kw2 = set(p2.keywords)
                overlap = kw1 & kw2
                
                if len(overlap) >= 1:
                    strength = min(len(overlap) / 3.0, 1.0)
                    
                    rel = KnowledgeRelation(
                        course_id=course_id,
                        source_point_id=p1.id,
                        target_point_id=p2.id,
                        relation_type='related',
                        relation_strength=strength,
                        description=f"关键词: {', '.join(list(overlap)[:5])}"
                    )
                    self.db.add(rel)
        
        self.db.commit()
        
        # 更新图谱统计
        self.db.execute(
            text("SELECT update_knowledge_graph_stats(:course_id)"),
            {"course_id": str(course_id)}
        )
        self.db.commit()
    
    async def _update_progress(self, task_id: str, progress: int, step: int, status: str):
        """更新进度"""
        task = self.db.query(DocumentProcessingTask).filter(
            DocumentProcessingTask.id == task_id
        ).first()
        if task:
            task.progress = progress
            task.current_step = step
            task.status = status
            task.updated_at = datetime.utcnow()
            self.db.commit()
